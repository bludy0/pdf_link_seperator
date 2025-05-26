import pdfplumber
import re
from collections import Counter
import docx
from urllib.parse import urlparse
import os

def clean_text(text):
    """XML uyumlu olmayan karakterleri ve fazla boşlukları temizler."""
    if not text:
        return text
    # Kontrol karakterlerini kaldır
    text = ''.join(c for c in text if c.isprintable() and ord(c) >= 32)
    return re.sub(r'\s+', '', text)  # Tüm boşlukları kaldır

def clean_text_for_word(text):
    """Word dökümanı için XML uyumlu metin temizleme"""
    if not text:
        return text
    # XML uyumsuz karakterleri kaldır
    text = ''.join(char for char in text if ord(char) >= 32 or char in '\n\r\t')
    # Unicode boşlukları normal boşluğa çevir
    text = ' '.join(text.split())
    return text

def is_valid_url(url):
    """URL'nin geçerli olup olmadığını kontrol eder."""
    try:
        parsed = urlparse(url)
        return bool(parsed.scheme and parsed.netloc)  # Scheme (http/https) ve netloc (domain) olmalı
    except:
        return False

def is_valid_doi(url):
    """DOI URL'sinin geçerli olup olmadığını kontrol eder."""
    try:
        # DOI formatı kontrol: https://doi.org/10.XXXX/YYYY şeklinde olmalı
        doi_format = re.compile(
            r'https?://doi\.org/10\.\d{4,}/[-\w.;()/]+[-\w]+$',
            re.IGNORECASE
        )
        return bool(doi_format.match(url)) and len(url) > 30
    except:
        return False

def extract_links_from_pdf(pdf_path):
    links = []
    # DOI linkleri için pattern
    doi_pattern = re.compile(
        r'https?://doi\.org/10\.\d{4,}/\S+?(?=["\s]|$)',
        re.IGNORECASE
    )
    
    # Genel URL'ler için pattern
    url_pattern = re.compile(
        r'(?:https?://|www\.)[a-zA-Z0-9\-\.]+\.[a-zA-Z]{2,}(?:/[^"\s<>]*)?',
        re.IGNORECASE
    )
    
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                # DOI linklerini bul
                doi_matches = doi_pattern.finditer(text)
                for match in doi_matches:
                    link = match.group().rstrip('.')
                    if len(link) > 30:  # Minimum DOI uzunluğu kontrolü
                        links.append(link)
                
                # Normal URL'leri bul
                url_matches = url_pattern.finditer(text)
                for match in url_matches:
                    link = match.group().rstrip('.')
                    # www. ile başlıyorsa https:// ekle
                    if link.startswith('www.'):
                        link = 'https://' + link
                    
                    # doi.org içermeyen URL'leri ekle
                    if 'doi.org' not in link and len(link) > 10:
                        links.append(link)
            
            # PDF annotasyonlarından URL'leri al
            for annot in (page.annots or []):
                if annot.get('uri'):
                    uri = annot['uri'].strip()
                    if len(uri) > 10:
                        links.append(uri)
    
    return links

def get_domain(url):
    try:
        parsed = urlparse(url)
        return parsed.netloc or parsed.path.split('/')[0]
    except:
        return url

def create_frequency_document(links, output_path):
    link_counts = Counter(links)
    domain_counts = Counter(get_domain(link) for link in links)
    
    doc = docx.Document()
    doc.add_heading('PDF Link Frequency Analysis', 0)
    
    doc.add_heading('Links by Frequency', level=1)
    for link, count in link_counts.most_common():
        if link:  # Boş linkleri atla
            try:
                clean_link = clean_text_for_word(link)
                doc.add_paragraph(f'{clean_link}: {count} times')
            except Exception as e:
                print(f"Skipping problematic link: {link[:50]}... ({str(e)})")
        
    doc.add_heading('Domains by Frequency', level=1)
    for domain, count in domain_counts.most_common():
        if domain:  # Boş domainleri atla
            try:
                clean_domain = clean_text_for_word(domain)
                doc.add_paragraph(f'{clean_domain}: {count} times')
            except Exception as e:
                print(f"Skipping problematic domain: {domain[:50]}... ({str(e)})")
    
    # Çıktı dosyasının dizininin var olduğunu kontrol et
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    doc.save(output_path)

def main():
    # Specify the single PDF file path\
    ########################################################
    pdf_path = 'example.pdf' #'C:\\Users\\example.pdf'
    output_doc = 'example.docx' #'C:\\Users\\example.docx'
    ########################################################
    # Check if the PDF exists before processing
    if not os.path.exists(pdf_path):
        print(f"Error: PDF file '{pdf_path}' not found.")
        return  # Exit the function if the file is not found
    
    links = extract_links_from_pdf(pdf_path)
    
    if links:
        create_frequency_document(links, output_doc)
        print(f"Report generated: {output_doc}")
    else:
        print(f"No links found in the PDF: {pdf_path}")

if __name__ == "__main__":
    main()
